# app/services/document_processor.py
import os
import uuid
from typing import List, Tuple
import PyPDF2
from docx import Document as DocxDocument
from app.models.document import Document, DocumentChunk
from app.core.config import settings

class DocumentProcessor:
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
    
    async def process_file(self, file_path: str, filename: str, file_type: str) -> Document:
        """Process uploaded file and extract text"""
        text_content = await self._extract_text(file_path, file_type)
        chunks = self._create_chunks(text_content)
        
        file_size = os.path.getsize(file_path)
        
        document = Document(
            id=str(uuid.uuid4()),
            name=filename,
            file_type=file_type,
            size=file_size,
            upload_date=datetime.now(),
            text_content=text_content,
            chunks=chunks
        )
        
        return document
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from different file types"""
        if file_type == "application/pdf":
            return await self._extract_pdf_text(file_path)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            return await self._extract_docx_text(file_path)
        elif file_type in ["text/plain", "text/markdown"]:
            return await self._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    async def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _create_chunks(self, text: str) -> List[DocumentChunk]:
        """Split text into chunks"""
        chunks = []
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunk_text = text[start:end]
            
            chunk = DocumentChunk(
                id=f"chunk_{chunk_id}",
                text=chunk_text,
                start_index=start,
                end_index=end,
                metadata={"chunk_index": chunk_id}
            )
            
            chunks.append(chunk)
            chunk_id += 1
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break
        
        return chunks
